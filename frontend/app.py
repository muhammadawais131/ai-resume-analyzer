import streamlit as st
import requests
import io
from fpdf import FPDF
from docx import Document

# Backend URL
BACKEND_URL = "http://127.0.0.1:8000"

# Page Config
st.set_page_config(
    page_title="AI Resume Analyzer", 
    page_icon="📄", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Professional Modern CSS - Dark Blue Theme
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    /* Main App Background - Professional Dark Blue */
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Main Title */
    h1 {
        color: #ffffff;
        font-weight: 700;
        font-size: 2.5rem;
        text-align: center;
        padding: 2rem 0;
        margin-bottom: 2rem;
        text-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    
    /* Section Headers */
    h3 {
        color: #60a5fa;
        font-weight: 600;
        font-size: 1.5rem;
        margin: 2rem 0 1rem 0;
    }
    
    /* Form Container */
    .stForm {
        background: rgba(255, 255, 255, 0.95);
        padding: 2.5rem;
        border-radius: 16px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    
    /* File Uploader */
    .stFileUploader {
        background: #f1f5f9;
        border: 2px dashed #3b82f6;
        border-radius: 12px;
        padding: 2rem;
    }
    
    .stFileUploader label {
        color: #1e293b !important;
        font-weight: 600;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzone"] {
        color: #1e293b !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzone"] p {
        color: #1e293b !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzone"] small {
        color: #475569 !important;
    }
    
    /* Uploaded file name */
    .stFileUploader section[data-testid="stFileUploaderFileData"] {
        color: #1e293b !important;
    }
    
    .stFileUploader section[data-testid="stFileUploaderFileData"] span {
        color: #1e293b !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderFileName"] {
        color: #1e293b !important;
    }
    
    /* All text inside file uploader */
    .stFileUploader * {
        color: #1e293b !important;
    }
    
    .stFileUploader button {
        color: #3b82f6 !important;
    }
    
    /* Text Area */
    .stTextArea textarea {
        border: 2px solid #e2e8f0;
        border-radius: 10px;
        background: #f8fafc;
        color: #1e293b !important;
        font-size: 15px;
        padding: 1rem;
        caret-color: #1e293b !important;
        cursor: text;
    }
    
    .stTextArea textarea:focus {
        border-color: #3b82f6 !important;
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
        caret-color: #1e293b !important;
    }
    
    .stTextArea textarea::placeholder {
        color: #94a3b8 !important;
    }
    
    .stTextArea label {
        color: #1e293b !important;
        font-weight: 600;
        font-size: 16px;
    }
    
    .stTextArea div[data-baseweb="base-input"] {
        color: #1e293b !important;
    }
    
    /* Ensure cursor is visible */
    @keyframes blink-caret {
        from, to { border-color: transparent; }
        50% { border-color: #1e293b; }
    }
    
    /* Primary Button */
    .stButton button {
        background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
        color: white !important;
        border: none;
        padding: 1rem 3rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 18px;
        width: 100%;
        box-shadow: 0 4px 16px rgba(59, 130, 246, 0.4);
        transition: all 0.3s ease;
    }
    
    .stButton button:hover {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
        box-shadow: 0 6px 24px rgba(59, 130, 246, 0.6);
        transform: translateY(-2px);
    }
    
    /* Download Buttons */
    .stDownloadButton button {
        background: white;
        color: #3b82f6 !important;
        border: 2px solid #3b82f6;
        padding: 0.875rem 2rem;
        border-radius: 10px;
        font-weight: 600;
        font-size: 16px;
        transition: all 0.3s ease;
        width: 100%;
    }
    
    .stDownloadButton button:hover {
        background: #3b82f6;
        color: white !important;
        box-shadow: 0 4px 16px rgba(59, 130, 246, 0.4);
        transform: translateY(-2px);
    }
    
    /* Success/Info/Warning Messages */
    .stSuccess {
        background: rgba(16, 185, 129, 0.15);
        border-left: 4px solid #10b981;
        border-radius: 8px;
        padding: 1rem;
        color: #ffffff !important;
    }
    
    .stSuccess > div {
        color: #ffffff !important;
    }
    
    .stInfo {
        background: rgba(59, 130, 246, 0.15);
        border-left: 4px solid #3b82f6;
        border-radius: 8px;
        padding: 1rem;
        color: #ffffff !important;
    }
    
    .stInfo > div {
        color: #ffffff !important;
    }
    
    .stWarning {
        background: rgba(245, 158, 11, 0.15);
        border-left: 4px solid #f59e0b;
        border-radius: 8px;
        padding: 1rem;
        color: #ffffff !important;
    }
    
    .stWarning > div {
        color: #ffffff !important;
    }
    
    .stError {
        background: rgba(239, 68, 68, 0.15);
        border-left: 4px solid #ef4444;
        border-radius: 8px;
        padding: 1rem;
        color: #ffffff !important;
    }
    
    .stError > div {
        color: #ffffff !important;
    }
    
    /* Match Score Box - Attractive */
    .match-score-container {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        padding: 3rem;
        border-radius: 20px;
        text-align: center;
        box-shadow: 0 10px 40px rgba(59, 130, 246, 0.5);
        margin: 2rem 0;
        border: 1px solid rgba(255, 255, 255, 0.2);
    }
    
    .match-score-container h2 {
        color: #ffffff;
        font-size: 1.5rem;
        font-weight: 600;
        margin: 0 0 1rem 0;
        opacity: 0.95;
    }
    
    .match-score-container .score {
        color: #ffffff;
        font-size: 4.5rem;
        font-weight: 800;
        margin: 0.5rem 0;
        text-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
    }
    
    .match-score-container p {
        color: #ffffff;
        font-size: 1.1rem;
        margin: 0;
        opacity: 0.9;
    }
    
    /* Suggestions Container */
    .suggestions-container {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 16px;
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
    }
    
    .suggestions-container h3 {
        color: #1e293b !important;
        margin-top: 0;
    }
    
    .suggestions-container ul {
        list-style: none;
        padding: 0;
        margin: 0;
    }
    
    .suggestions-container li {
        background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
        border-left: 4px solid #3b82f6;
        padding: 1.25rem;
        margin: 1rem 0;
        border-radius: 10px;
        color: #1e293b !important;
        font-size: 15px;
        line-height: 1.7;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
    }
    
    /* Missing Skills Container */
    .missing-skills-container {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 16px;
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
    }
    
    .missing-skills-container h3 {
        color: #1e293b !important;
        margin-top: 0;
    }
    
    .missing-skills-container ul {
        list-style: none;
        padding: 0;
        display: flex;
        flex-wrap: wrap;
        gap: 0.75rem;
    }
    
    .missing-skills-container li {
        background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%);
        color: #dc2626 !important;
        padding: 0.75rem 1.5rem;
        border-radius: 25px;
        font-weight: 600;
        font-size: 14px;
        border: 2px solid #fca5a5;
        box-shadow: 0 2px 8px rgba(220, 38, 38, 0.15);
        transition: all 0.3s ease;
    }
    
    .missing-skills-container li:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(220, 38, 38, 0.25);
    }
    
    /* Preview Container */
    .preview-container {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 16px;
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
    }
    
    .preview-container h3 {
        color: #1e293b !important;
        margin-top: 0;
    }
    
    .preview-container textarea {
        color: #1e293b !important;
    }
    
    .preview-container label {
        color: #1e293b !important;
    }
    
    /* Export Section */
    .export-section {
        background: rgba(255, 255, 255, 0.95);
        border-radius: 16px;
        padding: 2.5rem;
        margin: 2rem 0;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
    }
    
    .export-section h3 {
        color: #1e293b !important;
        margin-top: 0;
        font-size: 1.75rem;
    }
    
    .export-section p {
        color: #1e293b !important;
    }
    
    .export-section div {
        color: #1e293b !important;
    }
    
    /* Divider */
    hr {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, rgba(59, 130, 246, 0.5), transparent);
        margin: 3rem 0;
    }
    
    /* Spinner */
    .stSpinner > div {
        border-top-color: #3b82f6;
    }
    
    /* Text in preview */
    .stTextArea [data-baseweb="textarea"] {
        background: #f8fafc;
        border: 2px solid #cbd5e1;
        color: #1e293b !important;
        caret-color: #1e293b !important;
        cursor: text;
    }
    
    .stTextArea [data-baseweb="textarea"]:focus {
        border-color: #3b82f6 !important;
        outline: none !important;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
        caret-color: #1e293b !important;
    }
    
    /* Form submit button text */
    .stFormSubmitButton button {
        color: white !important;
    }
    
    /* Ensure all paragraph text in white containers is dark */
    .suggestions-container p,
    .missing-skills-container p,
    .preview-container p,
    .export-section p {
        color: #1e293b !important;
    }
</style>
""", unsafe_allow_html=True)

st.title("📄 AI Resume Analyzer & Optimizer")

# Initialize session state
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'uploaded_file_data' not in st.session_state:
    st.session_state.uploaded_file_data = None
if 'job_description_data' not in st.session_state:
    st.session_state.job_description_data = None
if 'rewritten_resume' not in st.session_state:
    st.session_state.rewritten_resume = None

# ---- Resume Analysis Section ----
with st.form("resume_form"):
    uploaded_file = st.file_uploader("📄 Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
    job_description = st.text_area("💼 Paste Job Description Here", height=180, placeholder="Paste the complete job description here...")
    submitted = st.form_submit_button("🚀 Analyze Resume")

if submitted and uploaded_file and job_description:
    st.session_state.uploaded_file_data = {
        'name': uploaded_file.name,
        'type': uploaded_file.type,
        'content': uploaded_file.getvalue()
    }
    st.session_state.job_description_data = job_description
    
    with st.spinner("Analyzing resume... please wait"):
        files = {"file": (uploaded_file.name, uploaded_file, uploaded_file.type)}
        data = {"job_description": job_description}
        try:
            res = requests.post(f"{BACKEND_URL}/analyze", files=files, data=data, timeout=120)
            if res.status_code == 200:
                result = res.json()
                st.session_state.analysis_result = result
                st.success(result.get("message", "✅ Analysis completed successfully!"))

                # Match Score Display
                st.markdown(f"""
                <div class="match-score-container">
                    <h2>Match Score</h2>
                    <div class="score">{result['match_score']}%</div>
                    <p>Resume alignment with job requirements</p>
                </div>
                """, unsafe_allow_html=True)

                # Professional Suggestions
                st.markdown('<div class="suggestions-container">', unsafe_allow_html=True)
                st.markdown("### 💡Suggestions")
                if result.get("suggestions"):
                    suggestions_html = '<ul>'
                    for suggestion in result["suggestions"]:
                        suggestions_html += f"<li>{suggestion}</li>"
                    suggestions_html += "</ul>"
                    st.markdown(suggestions_html, unsafe_allow_html=True)
                else:
                    st.info("No suggestions found — great job!")
                st.markdown('</div>', unsafe_allow_html=True)

                # Missing Skills Section
                st.markdown('<div class="missing-skills-container">', unsafe_allow_html=True)
                st.markdown("### ❌ Missing Skills")
                if result.get("missing_skills"):
                    skills_html = '<ul>'
                    for skill in result["missing_skills"]:
                        skills_html += f"<li>{skill}</li>"
                    skills_html += "</ul>"
                    st.markdown(skills_html, unsafe_allow_html=True)
                else:
                    st.success("✅ All key skills covered!")
                st.markdown('</div>', unsafe_allow_html=True)

                # Get rewritten resume
                with st.spinner("Generating improved resume..."):
                    rewrite_files = {"file": (st.session_state.uploaded_file_data['name'], 
                                             io.BytesIO(st.session_state.uploaded_file_data['content']), 
                                             st.session_state.uploaded_file_data['type'])}
                    rewrite_data = {"job_description": st.session_state.job_description_data}
                    try:
                        rewrite_res = requests.post(f"{BACKEND_URL}/rewrite", 
                                                   files=rewrite_files, 
                                                   data=rewrite_data, 
                                                   timeout=120)
                        if rewrite_res.status_code == 200:
                            rewrite_result = rewrite_res.json()
                            rewritten_text = rewrite_result.get("rewritten_resume", "")
                            
                            # Clean text
                            rewritten_text = rewritten_text.replace("", "")
                            rewritten_text = rewritten_text.replace("*", "")
                            rewritten_text = rewritten_text.replace("###", "")
                            rewritten_text = rewritten_text.replace("##", "")
                            rewritten_text = rewritten_text.replace("#", "")
                            
                            st.session_state.rewritten_resume = rewritten_text
                        else:
                            st.warning("⚠ Could not generate rewritten resume.")
                            st.session_state.rewritten_resume = None
                    except requests.exceptions.RequestException as e:
                        st.warning(f"⚠ Rewrite service unavailable.")
                        st.session_state.rewritten_resume = None

                # Resume Preview
                if st.session_state.rewritten_resume:
                    st.markdown('<div class="preview-container">', unsafe_allow_html=True)
                    st.markdown("### 📝 AI-Rewritten Resume Preview")
                    st.text_area("Preview", st.session_state.rewritten_resume, height=300, key="preview_area")
                    st.markdown('</div>', unsafe_allow_html=True)

            else:
                st.error(f"Error: {res.text}")
        except requests.exceptions.RequestException as e:
            st.error(f"🚫 Error contacting backend: {e}")

# ---- Export Section ----
st.divider()

st.markdown('<div class="export-section">', unsafe_allow_html=True)
st.markdown("### 📤 Export Improved Resume")

if st.session_state.rewritten_resume:
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Export as PDF", key="export_pdf_btn"):
            with st.spinner("Generating PDF..."):
                try:
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.set_margins(15, 15, 15)
                    
                    resume_text = st.session_state.rewritten_resume
                    resume_text = resume_text.replace('•', '-').replace('●', '-').replace('○', '-').replace('■', '-').replace('?', '-')
                    
                    sections = resume_text.split('\n\n')
                    
                    for section in sections:
                        if not section.strip():
                            continue
                            
                        lines = section.split('\n')
                        first_line = lines[0].strip()
                        is_heading = (
                            first_line.isupper() or 
                            any(first_line.upper().startswith(h) for h in 
                                ['SUMMARY', 'WORK EXPERIENCE', 'EXPERIENCE', 'SKILLS', 
                                 'EDUCATION', 'PROFESSIONAL', 'CONTACT', 'PROJECTS'])
                        )
                        
                        if is_heading:
                            pdf.set_font("Arial", 'B', 14)
                            clean_header = first_line.encode('ascii', 'ignore').decode('ascii')
                            pdf.cell(0, 10, clean_header, ln=True)
                            pdf.ln(2)
                            
                            pdf.set_font("Arial", size=10)
                            for line in lines[1:]:
                                if line.strip():
                                    clean_line = line.strip()
                                    if clean_line.startswith('-'):
                                        clean_line = '  - ' + clean_line.lstrip('-').strip()
                                    clean_line = clean_line.encode('ascii', 'ignore').decode('ascii')
                                    pdf.multi_cell(0, 5, clean_line)
                        else:
                            pdf.set_font("Arial", size=10)
                            for line in lines:
                                if line.strip():
                                    clean_line = line.strip()
                                    if clean_line.startswith('-'):
                                        clean_line = '  - ' + clean_line.lstrip('-').strip()
                                    clean_line = clean_line.encode('ascii', 'ignore').decode('ascii')
                                    pdf.multi_cell(0, 5, clean_line)
                        
                        pdf.ln(3)
                    
                    pdf_output = pdf.output(dest='S').encode('latin-1')
                    
                    st.success("✅ PDF generated successfully!")
                    st.download_button(
                        "⬇ Download PDF", 
                        data=pdf_output, 
                        file_name="improved_resume.pdf",
                        mime="application/pdf",
                        key="download_pdf_btn"
                    )
                except Exception as e:
                    st.error(f"❌ Failed to generate PDF: {e}")
    
    with col2:
        if st.button("📄 Export as DOCX", key="export_docx_btn"):
            with st.spinner("Generating DOCX..."):
                try:
                    doc = Document()
                    
                    resume_text = st.session_state.rewritten_resume
                    parts = [p.strip() for p in resume_text.split("\n\n") if p.strip()]
                    
                    for p in parts:
                        if any(header in p.lower()[:30] for header in ["summary", "work experience", "skills", "education", "professional summary"]):
                            lines = p.split("\n")
                            doc.add_heading(lines[0].strip(), level=2)
                            for line in lines[1:]:
                                if line.strip():
                                    if line.strip().startswith("-") or line.strip().startswith("•"):
                                        doc.add_paragraph(line.strip().lstrip("-•").strip(), style='List Bullet')
                                    else:
                                        doc.add_paragraph(line.strip())
                        else:
                            for line in p.split("\n"):
                                if line.strip():
                                    if line.strip().startswith("-") or line.strip().startswith("•"):
                                        doc.add_paragraph(line.strip().lstrip("-•").strip(), style='List Bullet')
                                    else:
                                        doc.add_paragraph(line.strip())
                    
                    docx_output = io.BytesIO()
                    doc.save(docx_output)
                    docx_output.seek(0)
                    
                    st.success("✅ DOCX generated successfully!")
                    st.download_button(
                        "⬇ Download DOCX", 
                        data=docx_output.getvalue(), 
                        file_name="improved_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_btn"
                    )
                except Exception as e:
                    st.error(f"❌ Failed to generate DOCX: {e}")
else:
    st.info("👆 Please analyze a resume first to enable export options.")

st.markdown('</div>', unsafe_allow_html=True)