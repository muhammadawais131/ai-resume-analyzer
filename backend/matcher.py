import os
import re
import time
import numpy as np
import openai
from dotenv import load_dotenv
from docx import Document

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

EMBED_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4o-mini"

STOPWORDS = {
    "a","an","the","and","or","in","on","at","of","for","to","with","by","as","is","are","was","were",
    "that","this","these","those","it","its","be","from","which","but","have","has","had","will","can",
    "i","we","you","he","she","they","their","them","our","us"
}

PHRASE_MAP = {
    "customers": "Demonstrate experience interacting with customers or improving customer satisfaction.",
    "leadership": "Highlight leadership experience: team lead, project lead, or mentoring roles.",
    "warehouse": "Show hands-on warehouse experience, inventory management, or logistics tasks.",
    "hardworking": "Illustrate work ethic with examples of reliability, punctuality, and outcomes.",
    "support": "Describe support responsibilities: customer support, technical support or admin support.",
    "adaptable": "Provide examples of adapting to new tools, teams, or responsibilities.",
    "team": "Include team-based achievements and collaborative results.",
    "python": "List Python projects, libraries used (e.g., pandas, numpy), and impact/metrics.",
    "javascript": "List JS frameworks used (e.g., React, Node) and specific accomplishments."
}

# ----------------------
def get_embedding(text):
    """Get embedding with timeout and error handling for speed"""
    try:
        
        if len(text) > 8000:
            text = text[:8000]
        
        response = openai.embeddings.create(
            input=text,
            model=EMBED_MODEL,
            timeout=10  
        )
        return response.data[0].embedding
    except Exception:
        return np.zeros(1536)  # fallback embedding if API fails
def cosine_similarity(a, b):
    a, b = np.array(a), np.array(b)
    denom = (np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)

def _clean_and_tokenize(text):
    text = text.lower()
    text = re.sub(r"[^\w\s]", " ", text)
    tokens = [t for t in text.split() if t and t not in STOPWORDS and len(t) > 1]
    return tokens

def generate_suggestions(missing_skills, resume_text, job_text):
    """Generate actionable suggestions based on actual missing skills"""
    suggestions = []
    
    #Categorize  a missing skills
    technical_skills = []
    soft_skills = []
    
    tech_keywords = ['php', 'python', 'java', 'javascript', 'laravel', 'react', 'vue', 'mysql', 
                     'postgresql', 'mongodb', 'aws', 'azure', 'docker', 'git', 'api', 'rest',
                     'html', 'css', 'node', 'typescript', 'redis', 'nginx', 'linux']
    
    soft_keywords = ['leadership', 'communication', 'teamwork', 'management', 'problem-solving',
                     'analytical', 'collaboration', 'organization']
    
    for skill in missing_skills:
        skill_lower = skill.lower()
        if any(tech in skill_lower for tech in tech_keywords):
            technical_skills.append(skill)
        elif any(soft in skill_lower for soft in soft_keywords):
            soft_skills.append(skill)
    
    #suggestion
    if technical_skills:
        tech_list = ', '.join(technical_skills[:5])
        suggestions.append(f"Add these technical skills to your Skills section: {tech_list}")
    
    
    if soft_skills:
        soft_list = ', '.join(soft_skills[:3])
        suggestions.append(f"Demonstrate {soft_list} with specific examples in your work experience")
    
    
    has_metrics = resume_text.count('%') >= 1 or any(str(i) in resume_text for i in range(10, 100))
    if not has_metrics:
        suggestions.append("Add quantifiable achievements with numbers (e.g., 'Improved performance by 30%', 'Managed team of 10')")
    
    
    action_verbs = ['led', 'managed', 'developed', 'improved', 'increased', 'reduced', 'achieved', 'optimized', 'implemented', 'designed', 'built']
    has_action_verbs = any(verb in resume_text.lower() for verb in action_verbs)
    if not has_action_verbs:
        suggestions.append("Start bullet points with strong action verbs: Developed, Implemented, Optimized, Led, Achieved")
    
    
    if len(missing_skills) > 6:
        suggestions.append("Review the job description carefully and incorporate relevant keywords naturally throughout your resume")
    
    
    if 'summary' not in resume_text.lower()[:300]:
        suggestions.append("Add a professional summary at the top highlighting your key skills and experience")
    
    
    if len(suggestions) < 2:
        suggestions.append("Tailor your resume content to match the job requirements more closely")
    
    return suggestions[:5]

# ---------------------
def analyze_resume(resume_text, job_text):
    
    try:
        resume_embed = get_embedding(resume_text)
        job_embed = get_embedding(job_text)
        match_score = cosine_similarity(resume_embed, job_embed)
    except Exception:
        match_score = 0.0

    match_percentage = round(match_score * 100, 2)
    
    # Extract the actual skills from job description
    job_lower = job_text.lower()
    resume_lower = resume_text.lower()
    
    
    # Technical Skills
    technical_skills = {
        # Programming Languages
        'php', 'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 'rust', 'swift', 'kotlin',
        # Frameworks
        'laravel', 'symfony', 'codeigniter', 'django', 'flask', 'react', 'vue', 'angular', 'nodejs', 'express',
        'nextjs', 'nuxt', 'vuejs', 'reactjs', 'jquery',
        # Databases
        'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'mssql', 'mariadb', 'elasticsearch',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'gitlab', 'github', 'bitbucket', 'terraform',
        'ansible', 'nginx', 'apache',
        # Tools & Technologies
        'git', 'svn', 'jira', 'confluence', 'postman', 'swagger', 'webpack', 'composer', 'npm', 'yarn',
        # Web Technologies
        'html', 'html5', 'css', 'css3', 'sass', 'scss', 'bootstrap', 'tailwind', 'ajax', 'json', 'xml', 'rest',
        'restful', 'api', 'graphql', 'websocket', 'oauth', 'jwt',
        # Testing
        'phpunit', 'jest', 'mocha', 'selenium', 'cypress',
        # Other
        'linux', 'ubuntu', 'centos', 'windows', 'macos', 'bash', 'shell', 'ci/cd', 'agile', 'scrum', 'microservices'
    }
    

    soft_skills = {
        'leadership', 'communication', 'teamwork', 'problem-solving', 'analytical', 'management',
        'collaboration', 'organization', 'time-management', 'critical-thinking', 'adaptability',
        'creativity', 'attention-to-detail', 'self-motivated', 'reliable', 'punctual'
    }
    
    
    not_skills = {
        'understanding', 'integrate', 'develop', 'infrastructure', 'experienced', 'analytical',
        'job', 'rds', 'cloud', 'third', 'applications', 'experience', 'seeking', 'ideal',
        'candidate', 'responsibilities', 'required', 'preferred', 'qualifications', 'years',
        'strong', 'excellent', 'good', 'knowledge', 'ability', 'skills', 'work', 'working',
        'including', 'implement', 'maintain', 'optimize', 'ensure', 'build', 'create',
        'design', 'write', 'collaborate', 'participate', 'troubleshoot', 'deploy'
    }
    
    # Finding a missing skills
    missing_technical = []
    missing_soft = []
    
    
    for skill in technical_skills:
        if skill in job_lower and skill not in resume_lower:
            missing_technical.append(skill)
    
    
    for skill in soft_skills:
        if skill in job_lower and skill not in resume_lower:
            missing_soft.append(skill)
    
    
    missing_skills = missing_technical[:8] + missing_soft[:3]
    
    
    if len(missing_skills) < 3:
        resume_tokens = set(_clean_and_tokenize(resume_text))
        job_tokens = set(_clean_and_tokenize(job_text))
        candidate_missing = list(job_tokens - resume_tokens)
        
        
        for word in candidate_missing[:20]:
            if (word not in not_skills and 
                word not in STOPWORDS and
                len(word) > 3 and
                not word.isdigit() and
                len(missing_skills) < 10):
                
                if any(keyword in job_text.lower() for keyword in [f'{word} experience', f'{word} knowledge', f'{word} proficiency']):
                    missing_skills.append(word)
    
    # Generating a  suggestions
    suggestions = generate_suggestions(missing_skills, resume_text, job_text)
    
    return match_percentage, suggestions, missing_skills[:10]

def skill_gap_analysis(resume_text, job_text):
    _, _, missing = analyze_resume(resume_text, job_text)
    return {"missing_skills": missing}

def rewrite_resume(resume_text, job_text=None):
    system_prompt = """You are a professional resume writer. Create a clean, professional resume without any markdown formatting, asterisks, or special characters. 
    
    Use these formatting rules:
    - Section headers should be in UPPERCASE (e.g., SUMMARY, WORK EXPERIENCE, SKILLS, EDUCATION)
    - Use ONLY simple dash character (-) for bullet points, NO special unicode bullets (•, ●, ○)
    - No bold (), italic (*), or any markdown formatting
    - No special characters, only basic ASCII
    - Write in clear, concise, professional language
    - Focus on achievements with metrics and results
    - Keep it ATS-friendly
    
    Create sections: SUMMARY, WORK EXPERIENCE, SKILLS, EDUCATION
    
    Example bullet format:
    - Achieved 25% increase in sales
    - Managed team of 10 employees"""
    
    user_prompt = f"Resume:\n{resume_text}\n\n"
    if job_text:
        user_prompt += f"Job Description:\n{job_text}\n\nRewrite this resume to be professional, tailored for this job, and formatted as specified."
    else:
        user_prompt += "Rewrite this resume to be professional and formatted as specified."

    try:
        response = openai.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=1500,
            temperature=0.2
        )
        rewritten = response.choices[0].message.content.strip()
        
        
        rewritten = rewritten.replace("", "").replace("*", "")
        rewritten = rewritten.replace('•', '-')
        rewritten = rewritten.replace('●', '-')
        rewritten = rewritten.replace('○', '-')
        rewritten = rewritten.replace('■', '-')
        rewritten = rewritten.replace('◆', '-')
        rewritten = rewritten.replace('▪', '-')
        
    except Exception as e:
        # If API fails,return a formatted version of original text
        rewritten = f"SUMMARY\n{resume_text[:500]}\n\n(Note: Automatic rewriter is temporarily unavailable. Please check your OPENAI_API_KEY configuration.)"
    
    return rewritten

def generate_docx(rewritten_text, output_path=None):
    if output_path is None:
        timestamp = int(time.time())
        output_path = f"improved_resume_{timestamp}.docx"

    doc = Document()
    parts = [p.strip() for p in rewritten_text.split("\n\n") if p.strip()]
    
    for p in parts:
        lines = p.split("\n")
        first_line = lines[0].strip()
        
        
        is_heading = (
            first_line.isupper() or 
            any(first_line.upper().startswith(h) for h in 
                ['SUMMARY', 'WORK EXPERIENCE', 'EXPERIENCE', 'SKILLS', 
                 'EDUCATION', 'PROFESSIONAL', 'CONTACT', 'PROJECTS'])
        )
        
        if is_heading:
            doc.add_heading(first_line, level=2)
            for line in lines[1:]:
                if line.strip():
                    if line.strip().startswith("-") or line.strip().startswith("•"):
                        doc.add_paragraph(line.strip().lstrip("-•").strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
        else:
            for line in lines:
                if line.strip():
                    if line.strip().startswith("-") or line.strip().startswith("•"):
                        doc.add_paragraph(line.strip().lstrip("-•").strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())
    
    doc.save(output_path)
    return output_path